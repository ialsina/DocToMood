import docx
from docx.enum.text import WD_COLOR_INDEX
import math
import pandas as pd
from xml.sax.saxutils import escape



def get_docx(filename, join=False):
    doc = docx.Document(filename)
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    if join:
        return "\n".join(paragraphs)
    return paragraphs

def get_docx_with_highlight_mark(filename):
    doc = docx.Document(filename)
    paragraphs = []

    for para in doc.paragraphs:
        full_text = para.text.strip()
        highlighted = False

        for run in para.runs:
            if run.font.highlight_color is not None and run.text.strip():
                highlighted = True
                break

        if highlighted:
            full_text += " [HIGHLIGHTED]"

        paragraphs.append(full_text)

    return paragraphs


def df_to_docx(df, output_path="questions.docx"):

    doc = docx.Document()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "question"
    hdr_cells[1].text = "A"
    hdr_cells[2].text = "B"
    hdr_cells[3].text = "C"
    hdr_cells[4].text = "D"
    hdr_cells[5].text = "extra"

    for _, row in df.iterrows():

        # safe read of correct value
        correct_val = row["correct"]

        # default: no highlight
        highlight_index = None

        # numeric? finite? 0–3?
        if (
            isinstance(correct_val, (int, float)) and 
            not math.isnan(correct_val)
        ):
            correct_int = int(correct_val)
            if 0 <= correct_int <= 3:
                highlight_index = 1 + correct_int  # A=1, …

        cells = table.add_row().cells
                                
        cells[0].text = str(row["question"])
        cells[1].text = str(row["ans0"])
        cells[2].text = str(row["ans1"])
        cells[3].text = str(row["ans2"])
        cells[4].text = str(row["ans3"])
        cells[5].text = str(row["extra"])

        # apply highlight if valid index
        if highlight_index is not None:
            cell_to_mark = cells[highlight_index]
            for paragraph in cell_to_mark.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(output_path)
    print(f"Saved to {output_path}")

def df_to_xml(df: pd.DataFrame, output_path="moodle_questions.xml"):

    def wrap_cdata(text):
        return f"<![CDATA[{text}]]>"

    xml = ['<quiz>']

    for i, row in df.iterrows():

        # define question name id
        qname = row.get("extra")
        if not isinstance(qname, str) or qname.strip() == "":
            qname = f"q_{i+1}"

        # question text html formatted
        question_html = f"<p><strong>{i+1}.</strong> {escape(str(row['question']))}</p>"

        xml.append('  <question type="multichoice">')
        xml.append('    <name>')
        xml.append(f"      <text>{escape(qname)}</text>")
        xml.append('    </name>')

        xml.append('    <questiontext format="html">')
        xml.append(f"      <text>{wrap_cdata(question_html)}</text>")
        xml.append('    </questiontext>')

        xml.append('    <generalfeedback format="html">')
        xml.append("      <text><![CDATA[]]></text>")
        xml.append('    </generalfeedback>')

        xml.append("    <defaultgrade>1.0000000</defaultgrade>")
        xml.append("    <penalty>0.3333333</penalty>")
        xml.append("    <hidden>0</hidden>")
        xml.append("    <single>true</single>")
        xml.append("    <shuffleanswers>true</shuffleanswers>")
        xml.append("    <answernumbering>abc</answernumbering>")

        answers = [row['ans0'], row['ans1'], row['ans2'], row['ans3']]
        correct_index = row['correct']

        for j, ans in enumerate(answers):
            if isinstance(correct_index, (int, float)) and int(correct_index) == j:
                fraction = "100"
            else:
                fraction = "0"

            xml.append(f'    <answer fraction="{fraction}" format="html">')
            xml.append(f"      <text>{wrap_cdata(str(ans))}</text>")
            xml.append('      <feedback format="html">')
            xml.append("        <text><![CDATA[]]></text>")
            xml.append("      </feedback>")
            xml.append("    </answer>")

        # feedback blocks
        xml.append('    <correctfeedback format="html">')
        xml.append('      <text><![CDATA[¡Correcto!]]></text>')
        xml.append('    </correctfeedback>')

        xml.append('    <partiallycorrectfeedback format="html">')
        xml.append('      <text><![CDATA[Parcialmente correcto.]]></text>')
        xml.append('    </partiallycorrectfeedback>')

        xml.append('    <incorrectfeedback format="html">')
        xml.append('      <text><![CDATA[Incorrecto. Revisa la explicación y vuelve a intentarlo.]]></text>')
        xml.append('    </incorrectfeedback>')

        xml.append("  </question>")

    xml.append("</quiz>")

    # write XML file
    xml_text = "\n".join(xml)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(xml_text)

    print(f"Saved Moodle XML to {output_path}")
    return xml_text
